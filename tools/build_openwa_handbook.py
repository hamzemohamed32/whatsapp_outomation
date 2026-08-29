from __future__ import annotations

from pathlib import Path
from datetime import date
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "artifacts"
WORK_DIR = OUT_DIR / ".handbook_work"
OUT_PATH = OUT_DIR / "OpenWA_v0.23.2_Operations_and_SaaS_Handbook.docx"
ARCH_PATH = WORK_DIR / "openwa_saas_architecture.png"
LOGO_SRC = ROOT / "docs" / "logo" / "openwa_logo.webp"
LOGO_PNG = WORK_DIR / "openwa_logo.png"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "1F6B45"
GREEN_FILL = "E8F4ED"
GOLD = "7A5A00"
GOLD_FILL = "FFF4D6"
RED = "9B1C1C"
RED_FILL = "FCE8E8"
WHITE = "FFFFFF"
BLACK = "161B22"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_font(run, *, size: float | None = None, bold: bool | None = None,
             italic: bool | None = None, color: str | None = None,
             name: str = "Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_single_column_geometry(table) -> None:
    """Apply the guide preset's exact 9360-DXA one-column table geometry."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_pr.append(layout)
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    table_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    col = OxmlElement("w:gridCol")
    col.set(qn("w:w"), "9360")
    grid.append(col)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        cell._tc.get_or_add_tcPr().append(tc_w)
    tc_w.set(qn("w:w"), "9360")
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def fixed_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
                widths: Sequence[float], *, header_fill: str = LIGHT_BLUE,
                font_size: float = 9.2) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    table_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_pr.append(layout)
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    table_pr.append(tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    width_dxa = [round(w * 1440) for w in widths]
    for w in width_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    all_rows = [headers, *rows]
    for r_idx, values in enumerate(all_rows):
        row = table.rows[0] if r_idx == 0 else table.add_row()
        prevent_row_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(values):
            cell = row.cells[c_idx]
            cell.width = Inches(widths[c_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa[c_idx]))
            tc_w.set(qn("w:type"), "dxa")
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            set_font(run, size=font_size, bold=(r_idx == 0), color=(NAVY if r_idx == 0 else BLACK))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def callout(doc: Document, title: str, text: str, *, kind: str = "note") -> None:
    palette = {
        "note": (LIGHT_BLUE, NAVY),
        "positive": (GREEN_FILL, GREEN),
        "caution": (GOLD_FILL, GOLD),
        "risk": (RED_FILL, RED),
    }
    fill, ink = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_single_column_geometry(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=120, end=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True, color=ink)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(text)
    set_font(r2, size=10, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), color)
    r_pr.append(r_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    txt = OxmlElement("w:t")
    txt.text = text
    run.append(txt)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, field_code: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, separate, fallback, end):
        run = OxmlElement("w:r")
        run.append(el)
        paragraph._p.append(run)


def make_num(doc: Document, *, bullet: bool, marker: str | None = None) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), (marker or "•") if bullet else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    # OOXML requires every abstractNum before the concrete num records. Inserting
    # an abstract after existing num nodes makes Word repair the list mapping and
    # can turn bullets into a continuation of the decimal list.
    first_num_index = next((i for i, child in enumerate(numbering) if child.tag == qn("w:num")), len(numbering))
    numbering.insert(first_num_index, abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def list_item(doc: Document, text: str, num_id: int, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True, size=11, color=BLACK)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2, size=11, color=BLACK)
    else:
        r = p.add_run(text)
        set_font(r, size=11, color=BLACK)


def body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        a = p.add_run(bold_lead)
        set_font(a, size=11, bold=True, color=BLACK)
        b = p.add_run(text[len(bold_lead):])
        set_font(b, size=11, color=BLACK)
    else:
        r = p.add_run(text)
        set_font(r, size=11, color=BLACK)


def code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_single_column_geometry(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    set_cell_margins(cell, top=120, bottom=120, start=120, end=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for i, line in enumerate(code.splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, size=8.6, name="Consolas", color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def section(doc: Document, number: int, title: str) -> None:
    # Put the page-break behavior on a tiny spacer paragraph. Word can place a
    # Heading 1 with pageBreakBefore into the top margin when a preceding table
    # nearly fills its page; the spacer gives the heading a stable text frame.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.page_break_before = True
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 1.0
    sr = spacer.add_run("\u200b")
    set_font(sr, size=1, color=WHITE)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.add_run(f"{number}. {title}")


def subhead(doc: Document, title: str, *, page_break_before: bool = False) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.keep_with_next = True
    p.add_run(title)


def create_logo() -> None:
    im = Image.open(LOGO_SRC).convert("RGBA")
    bg = Image.new("RGBA", im.size, (255, 255, 255, 0))
    bg.alpha_composite(im)
    bg.save(LOGO_PNG)


def load_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def create_architecture_diagram() -> None:
    w, h = 1500, 740
    im = Image.new("RGB", (w, h), "#F8FAFC")
    d = ImageDraw.Draw(im)
    title_font = load_font(42, True)
    head_font = load_font(27, True)
    body_font = load_font(22)
    small_font = load_font(18)
    d.text((70, 42), "Recommended commercial architecture", font=title_font, fill="#0B2545")
    d.text((70, 98), "Keep customer identity, consent, quotas and billing outside the gateway.", font=body_font, fill="#5F6B7A")

    boxes = [
        (70, 190, 310, 350, "Customer app", "Users • opt-in • templates"),
        (390, 190, 650, 350, "SaaS control plane", "Tenants • RBAC • billing"),
        (730, 190, 980, 350, "Queue & policy", "Pacing • quotas • dedupe"),
        (1060, 190, 1430, 350, "OpenWA deployment", "API • session • webhooks"),
    ]
    for x1, y1, x2, y2, title, note in boxes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=20, fill="#FFFFFF", outline="#2E74B5", width=4)
        d.text((x1 + 24, y1 + 28), title, font=head_font, fill="#0B2545")
        d.multiline_text((x1 + 24, y1 + 82), note, font=small_font, fill="#5F6B7A", spacing=8)
    for x in (310, 650, 980):
        d.line((x + 12, 270, x + 68, 270), fill="#2E74B5", width=5)
        d.polygon(((x + 68, 270), (x + 48, 257), (x + 48, 283)), fill="#2E74B5")

    d.rounded_rectangle((1060, 425, 1430, 625), radius=18, fill="#E8F4ED", outline="#1F6B45", width=4)
    d.text((1090, 455), "Dedicated WhatsApp number", font=head_font, fill="#1F6B45")
    d.text((1090, 505), "One isolated session or, preferably,\none deployment per paying customer", font=small_font, fill="#234B38", spacing=8)
    d.line((1245, 350, 1245, 410), fill="#1F6B45", width=5)
    d.polygon(((1245, 410), (1232, 390), (1258, 390)), fill="#1F6B45")
    d.rounded_rectangle((70, 425, 980, 625), radius=18, fill="#FFF4D6", outline="#7A5A00", width=4)
    d.text((100, 455), "Do not confuse multi-session with multi-tenant", font=head_font, fill="#7A5A00")
    d.multiline_text((100, 510), "OpenWA v0.23.2 does not implement tenant users, tenant data boundaries, billing or 2FA.\nBuild these in the control plane or isolate each customer in a separate deployment.", font=body_font, fill="#5D4600", spacing=10)
    im.save(ARCH_PATH, quality=95)


def configure_document(doc: Document) -> tuple[int, int]:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    sec.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("OPENWA v0.23.2  |  OPERATIONS & SAAS HANDBOOK")
    set_font(r, size=8.5, bold=True, color=MUTED)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("OpenWA handbook  •  ")
    set_font(r, size=8.5, color=MUTED)
    add_field(fp, "PAGE")
    first_header = sec.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = sec.first_page_footer
    first_footer.paragraphs[0].text = ""
    return make_num(doc, bullet=True), make_num(doc, bullet=False)


def build() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    WORK_DIR.mkdir(exist_ok=True)
    create_logo()
    create_architecture_diagram()
    doc = Document()
    bullet_id, decimal_id = configure_document(doc)
    props = doc.core_properties
    props.title = "OpenWA v0.23.2: Operations and SaaS Handbook"
    props.subject = "Capabilities, safe sending, plugins, webhooks, deployment, security, templates, and SaaS architecture"
    props.author = "OpenWA local deployment handbook"
    props.keywords = "OpenWA, WhatsApp, API, webhook, plugin, Docker, SaaS, security"
    props.comments = "Prepared from the local OpenWA v0.23.2 repository and official WhatsApp policy sources."

    # Editorial cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(22)
    pic = p.add_run().add_picture(str(LOGO_PNG), width=Inches(1.45))
    pic._inline.docPr.set("descr", "OpenWA logo")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("OPENWA v0.23.2")
    set_font(r, size=13, bold=True, color=GREEN)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Operations & SaaS Handbook")
    set_font(r, size=29, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(38)
    r = p.add_run("Capabilities, realistic sending limits, plugins, webhooks, deployment, security, and ready-to-use templates")
    set_font(r, size=13.5, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Prepared for the local Windows installation and future customization")
    set_font(r, size=10.5, italic=True, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Repository: rmyndharis/OpenWA  •  Version: 0.23.2  •  24 August 2026")
    set_font(r, size=9.5, color=MUTED)

    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run("How to use this handbook")
    set_font(r, size=20, bold=True, color=NAVY)
    body(doc, "Read Sections 1–3 first. They answer what the project can do, how fast it can send, and where the account-ban risk comes from. Use Sections 7–10 before exposing the system publicly or selling it as a service.")
    callout(doc, "The most important fact", "OpenWA is an unofficial WhatsApp Web gateway. It is powerful and fully self-hosted, but it is not Meta’s official WhatsApp Business Platform. No messages-per-second setting can eliminate restriction or ban risk.", kind="risk")
    subhead(doc, "Contents")
    contents = [
        "Executive answer", "What OpenWA is—and is not", "Messaging capacity and safe pacing",
        "Capability map", "Choosing the WhatsApp engine", "Dashboard and API workflow",
        "Plugins explained", "Webhooks explained", "Deployment guide", "Security hardening",
        "SaaS architecture and gaps", "Ready-to-use message templates", "Operations and upgrades",
        "Go-live checklist", "Glossary and sources",
    ]
    for item in contents:
        list_item(doc, item, decimal_id)

    section(doc, 1, "Executive answer")
    callout(doc, "How many messages per second?", "There is no guaranteed safe number. In the built-in bulk sender, one batch defaults to one message every 3–5 seconds—about 12–20 per minute, roughly 15/minute on average. The API permits a 1-second minimum delay, but that is only a technical setting. It is not a WhatsApp-approved or ban-safe rate.", kind="caution")
    fixed_table(doc,
        ["Question", "Clear answer"],
        [
            ["What can it do?", "Run multiple WhatsApp sessions; send/receive text and media; manage chats, groups, status, calls, labels and channels; expose REST, WebSocket, webhooks, dashboard, SDKs and an optional MCP surface."],
            ["Can I use all of it locally?", "Yes. The complete source is local, editable and buildable. Development uses Node 24 and ports 2785/2886; production is designed for Docker."],
            ["Are plugins like browser plugins?", "No. They are server-side extension packages that react to hooks, transform traffic, call allowed capabilities, store configuration, and integrate systems such as Chatwoot or Typebot."],
            ["What are webhooks?", "Outbound HTTPS notifications from OpenWA to your application when an event occurs—for example message.received or session.disconnected."],
            ["Is it already a SaaS platform?", "No. It is multi-session, not multi-tenant. Tenant users, isolation, billing, quotas, 2FA and self-service signup are not implemented."],
            ["Can it be commercial?", "Technically yes, but a commercial multi-customer service needs a separate control plane and strong isolation. For regulated, high-volume or mission-critical messaging, use Meta’s official WhatsApp Business Platform."],
        ], [1.65, 4.85])
    subhead(doc, "Workload fit")
    fixed_table(doc,
        ["Best fit", "Poor fit"],
        [
            ["Personal automation, internal tools, prototypes and learning.", "Cold marketing blasts, purchased lists or attempts to bypass enforcement."],
            ["Low-volume support and expected notifications on a dedicated number.", "Regulated or mission-critical processing dependent on an unofficial client."],
            ["Workloads with SMS, email or official-API fallback.", "Public multi-tenant SaaS without isolation, consent controls and incident response."],
        ], [3.25, 3.25], font_size=8.7)

    section(doc, 2, "What OpenWA is—and is not")
    body(doc, "OpenWA v0.23.2 is a self-hosted NestJS API gateway with a React dashboard. It connects through reverse-engineered clients—whatsapp-web.js or Baileys—rather than Meta’s Cloud API. You own the code, infrastructure, database, session credentials and operational risk.")
    fixed_table(doc,
        ["Area", "OpenWA v0.23.2", "Not included automatically"],
        [
            ["Access", "REST API, Swagger, dashboard, WebSocket events, JavaScript/Python/Go SDK surfaces", "A Meta-issued access token or official WhatsApp Business Account"],
            ["Accounts", "Many WhatsApp sessions in one deployment", "True tenant organizations, named users, tenant billing and tenant-level data boundaries"],
            ["Messages", "Text, media, reactions, edits, replies, status tracking and bulk batches", "Guaranteed delivery, official messaging tiers or a ban-safe throughput promise"],
            ["Automation", "Webhooks, hooks, plugins, integration fabric, n8n/community integrations", "A no-code business workflow tailored to your company"],
            ["Operations", "Docker, SQLite/PostgreSQL, Redis, MinIO/S3 backup, health probes, metrics", "Public TLS, domain name, cloud account, backups policy or 24/7 operations team"],
        ], [1.3, 2.65, 2.55], font_size=8.9)
    callout(doc, "Version clarity", "This repository is the current OpenWA gateway v0.23.2. It is not @open-wa/wa-automate v4.76 and it is not that library’s v5 alpha. Treat them as different products/codebases.", kind="note")
    body(doc, "Ownership is a strength: you can modify controllers, engines, dashboard pages, database adapters, plugins and deployment files. It also means you must maintain your fork, test upgrades and protect session data.")

    section(doc, 3, "Messaging capacity and safe pacing")
    subhead(doc, "Technical behavior in this repository")
    fixed_table(doc,
        ["Control", "Current behavior", "Meaning"],
        [
            ["Single sends", "Executed directly against the selected WhatsApp engine", "No BullMQ message queue is used for the hot send path."],
            ["Bulk request", "Up to 100 entries; exact duplicates collapse", "Each accepted batch is processed asynchronously."],
            ["Default bulk delay", "3,000 ms + random 0–2,000 ms", "One batch produces about 12–20 messages/minute; average is about 15/minute."],
            ["Minimum bulk delay", "1,000 ms", "At most about 1 message/second per batch when randomization is disabled."],
            ["Concurrent batches", "Default process cap: 50", "This is a memory guard, not permission to run 50 high-rate campaigns."],
            ["HTTP throttles", "Defaults include 10 requests/second burst, 100/minute, 1,000/hour", "Limits API calls—not WhatsApp’s account enforcement."],
            ["Send pacing", "Daily warm-up/cold-contact caps; OFF by default", "Enable it explicitly to refuse risky sends with HTTP 429 SEND_PACING_LIMITED."],
        ], [1.5, 2.2, 2.8], font_size=8.8)
    subhead(doc, "Recommended starting policy")
    policy_num = make_num(doc, bullet=False)
    for text in [
        "Enable SEND_PACING_ENABLED=true before real automation.",
        "Keep the supplied warm-up schedule (20→1000/day) and cold-contact schedule (5→100/day); prefer zero cold outreach.",
        "Use default bulk delay and randomization. For a new number, increase delay to 10–30 seconds and keep volumes well below the daily cap.",
        "Do not run parallel batches for the same session. Put a per-session queue in your SaaS control plane.",
        "Stop immediately when failure rates, blocks, reports, disconnections or restriction events rise.",
    ]:
        list_item(doc, text, policy_num)
    callout(doc, "No safe MPS guarantee", "The repository recommends a few messages per minute—not thousands/hour. Opt-in, recipient history, complaints, engine and IP reputation determine risk. Policy still requires opt-in, honored opt-outs and data protection; local templates are not Meta-approved.", kind="risk")

    section(doc, 4, "Capability map")
    fixed_table(doc,
        ["Domain", "Available capabilities", "Important notes"],
        [
            ["Sessions", "Create, start, stop, restart, logout, delete, QR/pairing, reconnect and per-session proxy", "Many sessions can run; resource cost depends heavily on engine."],
            ["Messaging", "Text, image, video, audio/voice note, document, sticker, contacts, location, reply, mentions, reactions, edit, revoke/delete, pin/star", "Engine parity is high but not perfect; unsupported routes return explicit errors."],
            ["Bulk & templates", "Bulk batches, progress/cancel, stored text templates with {{variables}}", "Bulk supports text/image/video/audio/document; templates are local text rendering."],
            ["Chats & contacts", "List/search chats and messages, read/unread, archive, mute, clear/delete, profile/status information", "Some operations need local message history on Baileys."],
            ["Groups", "Create/manage groups, invite links, participants/admins, settings and join requests", "Cold group additions consume the cold-reachout pacing budget when pacing is enabled."],
            ["Status, calls, presence", "Post/read/delete statuses, receive/reject calls, call events, presence updates", "Status recipient behavior differs by engine."],
            ["Channels & catalog", "Channel/newsletter operations, labels, catalog/product reads and selected sends", "Some channel media or catalog operations are engine-specific."],
            ["Automation", "Webhooks, WebSocket feed, plugin hooks, integration fabric, automation rules, n8n/community integrations", "Webhooks deliver outward; integration ingress receives provider webhooks inward."],
            ["Admin & observability", "RBAC API keys, audit operations, metrics, logs, health/readiness, queue dashboard, backup/import/export", "Message sends and webhook deliveries have dedicated records rather than ordinary audit-log rows."],
            ["AI/tool access", "Optional MCP server with 25 read-only tools or 51 tools when writes are enabled", "Keep read-only by default and issue a least-privilege session-scoped key."],
        ], [1.35, 3.0, 2.15], font_size=8.15)
    callout(doc, "Engine capability snapshot", "The repository’s generated matrix reports 112 interface methods across two adapters: 199 supported adapter cells and 25 unavailable cells. At the REST level, 90 methods are engine-neutral, 12 are Baileys-only and 9 are whatsapp-web.js-only.", kind="note")

    section(doc, 5, "Choosing the WhatsApp engine")
    fixed_table(doc,
        ["Decision", "whatsapp-web.js", "Baileys"],
        [
            ["Connection style", "Controls real WhatsApp Web in headless Chromium", "Uses the multi-device WebSocket protocol directly"],
            ["Approx. memory/session", "High: repository guidance estimates ~300–500 MB", "Low: repository guidance estimates ~30–80 MB"],
            ["Risk profile", "Lower relative fingerprint risk, but never zero", "Higher relative fingerprint risk"],
            ["Session density", "Fewer sessions per server", "Many more sessions per server"],
            ["Best use", "Dedicated high-value numbers where RAM is available", "Lab, internal use, higher density where the risk is accepted"],
            ["Recommendation", "Default choice when account safety matters most", "Use only after testing each required capability"],
        ], [1.35, 2.58, 2.57], font_size=8.8)
    body(doc, "Do not switch engines for an existing production session without a planned relink, regression test and rollback path. Engine differences include channel media, status recipients, chat-history-dependent actions, and several profile or moderation result signals.")

    section(doc, 6, "Dashboard and API workflow")
    subhead(doc, "Local development")
    code_block(doc, "npm ci\nnpm run dev\n\nDashboard: http://localhost:2886\nAPI:       http://localhost:2785/api\nSwagger:   http://localhost:2785/api/docs")
    dashboard_num = make_num(doc, bullet=False)
    for text in [
        "Open the dashboard and enter the API key generated/configured for the local instance.",
        "Create a session and choose whatsapp-web.js unless you deliberately want Baileys.",
        "Start the session, scan the QR code from WhatsApp → Linked devices, and wait for READY.",
        "Send one message to your own opted-in test number.",
        "Create a webhook and run its test action before enabling automation.",
    ]:
        list_item(doc, text, dashboard_num)
    subhead(doc, "Core API request")
    code_block(doc, "curl.exe -X POST http://localhost:2785/api/sessions/SESSION_ID/messages/send-text `\n  -H \"Content-Type: application/json\" `\n  -H \"X-API-Key: YOUR_API_KEY\" `\n  -d '{\"chatId\":\"254700000000@c.us\",\"text\":\"Hello from OpenWA\"}'")
    callout(doc, "API key safety", "The key is the dashboard login and API credential. Production keys should be long, rotated, role-limited, IP-limited where possible, and restricted to allowedSessions. Never put an ADMIN key in browser-side customer code.", kind="risk")

    section(doc, 7, "Plugins explained")
    body(doc, "A plugin is code loaded by the OpenWA server to extend behavior. It is not a Chrome extension and not a ChatGPT plugin. Official/community packages can connect OpenWA with products such as Chatwoot or Typebot, transform events, moderate outbound messages, auto-reply, add search providers, or receive provider webhooks through the integration fabric.")
    fixed_table(doc,
        ["Plugin concept", "Meaning in OpenWA", "Example"],
        [
            ["Manifest", "Declares identity, entry file, permissions, hooks, network allowlist and optional ingress routes", "A Chatwoot adapter declares only the network hosts and webhook route it needs."],
            ["Hook", "A lifecycle event delivered to plugin code", "message:received triggers routing; message:sending can modify or block an outbound send."],
            ["Capability", "A host-provided operation gated by manifest permission and session scope", "Send a reply, read session state or write plugin storage."],
            ["Per-session activation", "Plugin can be enabled/configured for selected WhatsApp sessions", "Customer A can use one integration while Customer B does not."],
            ["Ingress", "An external service calls an OpenWA-owned URL that dispatches to the plugin", "Chatwoot sends an event into the adapter."],
            ["Catalog/install", "Install a validated ZIP by upload, URL or configured catalog", "Production URL installs should carry a SHA-256 integrity pin."],
        ], [1.35, 2.75, 2.4], font_size=8.7)
    subhead(doc, "Safe plugin workflow")
    plugin_num = make_num(doc, bullet=False)
    for text in [
        "Review the source, manifest permissions, network allowlist, package hash and maintainer before installation.",
        "Install with an unscoped ADMIN key; then activate/configure only for the required sessions.",
        "Start with a test session and inspect logs, outbound calls and message behavior.",
        "Keep secrets in plugin configuration, not source control; rotate them after suspected exposure.",
        "Run genuinely untrusted plugins in a separate container/VM and connect over the OpenWA API.",
    ]:
        list_item(doc, text, plugin_num)
    callout(doc, "Security truth", "The worker-thread sandbox contains bugs and mediates declared capabilities, but it is not an OS security boundary. A malicious plugin shares the OpenWA process user’s filesystem and OS privileges. Install only code you trust.", kind="risk")

    section(doc, 8, "Webhooks explained")
    callout(doc, "Webhook in one sentence", "A webhook is an HTTPS callback: OpenWA sends a JSON POST when a subscribed event occurs; your server verifies the signature, deduplicates the delivery, returns 2xx quickly, and processes slow work asynchronously.")
    fixed_table(doc,
        ["Direction", "Mechanism", "Use"],
        [
            ["OpenWA → your app", "Webhook subscription", "Message, session, presence, group, status and call events"],
            ["Your app → OpenWA", "REST API", "Send messages and manage sessions/resources"],
            ["Provider → OpenWA plugin", "Integration Fabric ingress", "Receive Chatwoot/Typebot/provider events through a declared plugin route"],
            ["Dashboard/live client", "Socket.IO /events", "Real-time UI updates without polling"],
        ], [1.45, 2.0, 3.05], font_size=8.8)
    subhead(doc, "Available webhook events")
    body(doc, "Messages: received, sent, acknowledgement, failed, revoked, reaction and edited. Sessions: status, QR, authenticated, disconnected, reconnect loop and restriction. Other families: status received, presence update, group join/leave/update/join request, and call received/accepted/rejected/missed.")
    subhead(doc, "Create a signed webhook")
    code_block(doc, "curl.exe -X POST http://localhost:2785/api/sessions/SESSION_ID/webhooks `\n  -H \"Content-Type: application/json\" `\n  -H \"X-API-Key: YOUR_API_KEY\" `\n  -d '{\"url\":\"https://app.example.com/hooks/openwa\",\"events\":[\"message.received\",\"session.disconnected\"],\"secret\":\"REPLACE_WITH_32_RANDOM_CHARACTERS\",\"retryCount\":3}'")
    subhead(doc, "Receiver rules", page_break_before=True)
    receiver_num = make_num(doc, bullet=False)
    for text in [
        "Verify X-OpenWA-Signature as HMAC-SHA256 over the exact raw request body; the value begins sha256=.",
        "Deduplicate with X-OpenWA-Idempotency-Key or X-OpenWA-Delivery-Id.",
        "Return 2xx quickly. Put database work, AI calls and customer actions on your own queue.",
        "Keep WEBHOOK_SSRF_PROTECT=true. Explicitly allow internal targets only when necessary.",
        "Monitor delivery failures and redrive only after the receiver is healthy and idempotent.",
    ]:
        list_item(doc, text, receiver_num)
    fixed_table(doc,
        ["Delivery control", "Default/current behavior"],
        [
            ["Timeout", "10 seconds"],
            ["Retry delay", "5 seconds"],
            ["Retry attempts", "3 by default; configurable 0–5 per webhook"],
            ["Endpoints/session", "Default maximum 16"],
            ["Reliability", "Queue/outbox, recovery and dead-letter records are available; Redis queue is optional"],
            ["Filtering", "Optional AND conditions on sender, recipient, body, type, mentions, fromMe, hasMedia and isGroup"],
        ], [1.75, 4.75])

    section(doc, 9, "Deployment guide")
    subhead(doc, "Choose one deployment level")
    fixed_table(doc,
        ["Level", "Stack", "Use"],
        [
            ["Development", "npm run dev; SQLite/local; dashboard on 2886", "Code changes, debugging and one test session"],
            ["Basic production", "docker compose up -d; SQLite/local", "Single operator, low traffic, simple backup"],
            ["Production data", "docker compose --profile postgres up -d", "Durable database and easier operations"],
            ["Full stack", "docker compose --profile full up -d", "PostgreSQL + Redis + MinIO/S3-compatible storage"],
            ["Scale-out", "Kubernetes/Helm + PostgreSQL + Redis + shared storage", "Multiple replicas; requires session ownership and careful load routing"],
        ], [1.35, 2.65, 2.5], font_size=8.8)
    subhead(doc, "Recommended production sequence")
    deploy_num = make_num(doc, bullet=False)
    for text in [
        "Provision a Linux host; create production secrets outside source control and disable the development API key.",
        "Start PostgreSQL, then optional Redis/MinIO, then OpenWA.",
        "Place an HTTPS reverse proxy/load balancer in front; keep database, Redis, MinIO and Docker proxy private.",
        "Verify health probes, authentication, one session, one signed webhook and backup restore.",
        "Enable monitoring, automated backups and log retention before onboarding a real number.",
    ]:
        list_item(doc, text, deploy_num)
    code_block(doc, "docker compose up -d\ndocker compose --profile postgres up -d\ndocker compose --profile full up -d\ndocker compose ps\ndocker compose logs --tail 200 openwa-api")
    callout(doc, "Do not publish localhost development", "The Vite development dashboard on port 2886 and a development API key are for local work. Public production should expose only the HTTPS reverse proxy to port 2785 and should keep Swagger off unless deliberately enabled.", kind="risk")

    section(doc, 10, "Security hardening")
    fixed_table(doc,
        ["Priority", "Control", "Required action"],
        [
            ["P0", "Secrets", "Disable ALLOW_DEV_API_KEY; use long random API keys; set API_KEY_PEPPER; never commit .env or session data."],
            ["P0", "Access", "Use ADMIN only for lifecycle tasks; OPERATOR for sending; VIEWER for read-only; restrict allowedSessions and allowedIps."],
            ["P0", "Network", "HTTPS only; strict CORS_ORIGINS; firewall private services; keep WEBHOOK_SSRF_PROTECT=true."],
            ["P0", "Data", "Encrypt disks/backups; restrict database and S3 credentials; define retention and deletion; test restores."],
            ["P1", "API surface", "Keep ENABLE_SWAGGER=false in production; set BODY_SIZE_LIMIT; enable request throttles and per-session send pacing."],
            ["P1", "Webhooks", "Use 32+ random secret, verify raw-body HMAC, dedupe, rate-limit receiver and never log full payloads unnecessarily."],
            ["P1", "Plugins", "Trust review + SHA-256 pin + least permissions; OS-isolate untrusted code."],
            ["P1", "Docker", "Keep non-root execution; disable docker-proxy if dashboard datastore orchestration is unused."],
            ["P2", "Monitoring", "Alert on restriction/disconnect loops, send failures, webhook dead letters, auth failures, disk space and backup failures."],
            ["P2", "Updates", "Run tests and migrations in staging; read SECURITY.md; latest 0.23.x receives security fixes while older minors do not."],
        ], [0.55, 1.25, 4.7], font_size=8.2)
    subhead(doc, "Sensitive assets")
    for text in [
        "WhatsApp session credentials can control the linked account.",
        "API keys can read or change deployment data according to role and session scope.",
        "Messages, contacts and webhook payloads can contain personal data.",
        "Plugin configuration and custom webhook headers can contain third-party secrets.",
        "Docker orchestration access can become host-root-equivalent if the API container is compromised.",
    ]:
        list_item(doc, text, bullet_id)

    section(doc, 11, "SaaS architecture and gaps")
    body(doc, "The current release supports many sessions and session-scoped API keys, but its multitenancy document is explicitly a draft proposal. It does not implement tenant organizations, memberships, customer login, tenant 2FA, billing, quotas, per-tenant backups or a guaranteed tenant data boundary.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    pic = p.add_run().add_picture(str(ARCH_PATH), width=Inches(6.35))
    pic._inline.docPr.set("descr", "Recommended OpenWA commercial SaaS architecture showing customer app, SaaS control plane, queue and policy layer, OpenWA deployment, and dedicated WhatsApp number")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Figure 1. Recommended separation between SaaS control plane and OpenWA gateway")
    set_font(r, size=8.5, italic=True, color=MUTED)
    subhead(doc, "Recommended customer isolation")
    fixed_table(doc,
        ["Model", "Isolation", "Recommendation"],
        [
            ["One shared OpenWA instance", "Session-scoped keys only; no tenant entity", "Suitable only for one trusted organization or tightly controlled internal teams."],
            ["One OpenWA deployment/customer", "Separate process, DB, storage, secrets and network", "Recommended starting model for paying customers; simpler and safer tenant boundary."],
            ["Official WhatsApp Platform/customer", "Meta WABA/phone number and official policy controls", "Preferred for regulated, high-volume, revenue-critical or long-term commercial messaging."],
        ], [1.65, 2.35, 2.5], font_size=8.7)
    subhead(doc, "Your SaaS control plane must add", page_break_before=True)
    for text in [
        "Tenant organization, named users, password/SSO, 2FA, roles and support impersonation audit.",
        "Subscription/billing, trial rules, plan limits and customer lifecycle.",
        "Consent records: recipient, source, timestamp, categories, proof and opt-out state.",
        "Per-tenant/session queue, daily limits, cold-contact blocks, deduplication and idempotency.",
        "Template review, campaign approval, quiet hours, timezone handling and human escalation.",
        "Tenant-scoped database/storage/backup/export/deletion and cross-tenant security tests.",
        "Abuse detection, complaints, incident response, number replacement and alternate-channel fallback.",
    ]:
        list_item(doc, text, bullet_id)
    callout(doc, "Commercial recommendation", "Do not sell a shared-instance SaaS simply because the dashboard can create multiple sessions. Start with one isolated deployment per customer, or use the official WhatsApp Business Platform and build your SaaS control plane around it.", kind="caution")

    section(doc, 12, "Ready-to-use message templates")
    callout(doc, "Template meaning", "OpenWA templates are session-scoped text records with {{variable}} substitution, optional header and footer. They are not Meta-approved Business Platform templates and do not override opt-in, outreach or 24-hour-window rules.", kind="note")
    fixed_table(doc,
        ["Name", "Body", "Variables"],
        [
            ["order-confirmation", "Hi {{customer}}, we received order {{orderId}}. Total: {{total}}. We’ll update you when it ships.", "customer, orderId, total"],
            ["delivery-update", "Hi {{customer}}, order {{orderId}} is {{status}}. Track it here: {{trackingUrl}}", "customer, orderId, status, trackingUrl"],
            ["appointment-reminder", "Hi {{customer}}, reminder: {{service}} on {{date}} at {{time}}. Reply 1 to confirm or 2 to reschedule.", "customer, service, date, time"],
            ["support-received", "Hi {{customer}}, we received request {{ticketId}}. A human will reply by {{responseTime}}.", "customer, ticketId, responseTime"],
            ["payment-received", "Hi {{customer}}, payment {{paymentId}} for {{amount}} was received. Receipt: {{receiptUrl}}", "customer, paymentId, amount, receiptUrl"],
            ["otp-code", "Your verification code is {{code}}. It expires in {{minutes}} minutes. Do not share it.", "code, minutes"],
            ["opt-out-confirmed", "You have been unsubscribed from {{category}} messages. Reply START if you want to subscribe again.", "category"],
        ], [1.35, 3.85, 1.3], font_size=8.0)
    subhead(doc, "Create and send a local template")
    code_block(doc, "# Create\ncurl.exe -X POST http://localhost:2785/api/sessions/SESSION_ID/templates `\n  -H \"Content-Type: application/json\" -H \"X-API-Key: YOUR_API_KEY\" `\n  -d '{\"name\":\"order-confirmation\",\"body\":\"Hi {{customer}}, we received order {{orderId}}.\",\"footer\":\"Reply STOP to unsubscribe.\"}'\n\n# Render and send\ncurl.exe -X POST http://localhost:2785/api/sessions/SESSION_ID/messages/send-template `\n  -H \"Content-Type: application/json\" -H \"X-API-Key: YOUR_API_KEY\" `\n  -d '{\"chatId\":\"254700000000@c.us\",\"templateName\":\"order-confirmation\",\"variables\":{\"customer\":\"Amina\",\"orderId\":\"SO-1042\"}}'")
    subhead(doc, "Template rules", page_break_before=True)
    template_num = make_num(doc, bullet=False)
    for text in [
        "Use a unique name per session. Name limit is 100 characters; body 4,096; header/footer 1,024 each.",
        "Never place passwords, full payment-card details, national IDs or unnecessary sensitive data in templates.",
        "Validate every variable and URL before sending. Treat user-controlled template variables as untrusted input.",
        "Add a human escalation path and a clear opt-out for promotional or recurring traffic.",
        "Do not use the OTP template as the only authentication channel on this unofficial gateway; keep a fallback.",
    ]:
        list_item(doc, text, template_num)

    section(doc, 13, "Operations and upgrades")
    subhead(doc, "Daily checks")
    for text in [
        "All expected sessions are READY; investigate reconnect loops or restriction events immediately.",
        "Send failure and delivery/ack rates are stable; no unexpected volume spike or cold-contact pattern.",
        "Webhook receiver latency, retries and dead-letter count are normal.",
        "Disk, database, Redis, object storage and container health are normal.",
        "Backups completed and the most recent restore test is within policy.",
    ]:
        list_item(doc, text, bullet_id)
    subhead(doc, "Safe upgrade workflow for your local fork")
    upgrade_num = make_num(doc, bullet=False)
    for text in [
        "Commit or safely preserve your custom changes; never upgrade a dirty production checkout blindly.",
        "Fetch upstream and read CHANGELOG.md, SECURITY.md, migrations and engine capability changes.",
        "Merge/rebase on a dedicated codex/ or upgrade branch and resolve conflicts deliberately.",
        "Run npm ci, typecheck, lint, unit tests, dashboard tests, e2e tests and production build.",
        "Deploy to staging with a disposable WhatsApp number; test pairing, send/receive, webhooks, plugins, backup and restore.",
        "Back up production, deploy in a maintenance window, watch health/metrics/logs, and keep a rollback image and database plan.",
    ]:
        list_item(doc, text, upgrade_num)
    subhead(doc, "Failure response", page_break_before=True)
    fixed_table(doc,
        ["Signal", "Immediate action", "Do not"],
        [
            ["Restriction/ban warning", "Stop outbound automation; preserve logs; use WhatsApp appeal path; notify affected customer", "Keep retrying or rotate proxies to evade enforcement"],
            ["High send failures", "Open the pacing breaker; halt batches; verify session, recipient format and engine health", "Treat every 429 as a transient delivery failure"],
            ["Webhook backlog", "Pause dependent workflows; restore receiver; verify idempotency; redrive gradually", "Redrive all failures into a non-idempotent receiver"],
            ["Credential exposure", "Revoke/rotate API keys, webhook secrets and integrations; review audit and access logs", "Assume deleting a log line removes the risk"],
            ["Database/storage failure", "Stop writes if consistency is uncertain; restore tested backup; reconcile sessions and batches", "Resume interrupted bulk batches automatically and risk duplicates"],
        ], [1.25, 3.2, 2.05], font_size=8.25)

    section(doc, 14, "Go-live checklist")
    checks = [
        "Dedicated, non-critical WhatsApp number selected and warmed up.",
        "Recipient opt-in and opt-out evidence implemented.",
        "whatsapp-web.js/Baileys decision documented and required features tested.",
        "SEND_PACING_ENABLED=true and one queue per session enforced.",
        "Production API keys are least-privilege, session/IP scoped and rotated; development key disabled.",
        "HTTPS, firewall, strict CORS, body limit and request throttles configured.",
        "Swagger disabled unless deliberately required.",
        "Signed webhook verified against the exact raw body; receiver is idempotent.",
        "Plugins reviewed, hash-pinned and tested; untrusted code isolated outside the process.",
        "PostgreSQL/Redis/storage credentials private; backup and restore tested.",
        "Monitoring covers health, restrictions, failures, webhook backlog, auth failures and capacity.",
        "SaaS customers isolated; no shared-instance tenant claims without implemented tenant controls.",
        "Privacy policy, data retention/deletion and incident response approved for the operating jurisdiction.",
        "SMS/email/official-API fallback exists for critical workflows.",
        "Staging test and rollback drill completed.",
    ]
    check_num = make_num(doc, bullet=True, marker="☐")
    for check in checks:
        list_item(doc, check, check_num)
    callout(doc, "Launch gate", "If any P0 security item, consent control, tenant boundary, backup restore, or fallback path is missing, keep the deployment private and in testing.", kind="risk")

    section(doc, 15, "Glossary and sources")
    fixed_table(doc,
        ["Term", "Meaning"],
        [
            ["API key", "Secret credential used for dashboard and REST/MCP authentication."],
            ["Session", "One linked WhatsApp account and its engine state."],
            ["Engine", "The adapter that connects to WhatsApp: whatsapp-web.js or Baileys."],
            ["Webhook", "Outbound JSON event callback from OpenWA to another server."],
            ["Plugin", "Server-side code package extending OpenWA through hooks and capabilities."],
            ["Integration ingress", "Inbound provider webhook routed by the OpenWA host to a declared plugin route."],
            ["HMAC", "Keyed signature used to prove webhook body integrity and authenticity."],
            ["Idempotency", "Processing a repeated delivery without creating a duplicate effect."],
            ["Multi-session", "Several WhatsApp accounts in one deployment."],
            ["Multi-tenant", "Several independent customer organizations with enforced identity and data boundaries."],
            ["WABA", "WhatsApp Business Account used by Meta’s official Business Platform."],
        ], [1.55, 4.95], font_size=8.8)
    subhead(doc, "Repository sources reviewed", page_break_before=True)
    for source in [
        "README.md — product scope, risk warnings, features, setup and production profiles.",
        "SECURITY.md and docs/04-security-design.md — roles, session scope, SSRF, CORS, plugins and Docker risk.",
        "docs/10-devops-infrastructure.md and docs/13-horizontal-scaling.md — deployment and scaling.",
        "docs/19-plugin-architecture.md and docs/30-plugin-sandboxing.md — plugin lifecycle, permissions and trust model.",
        "docs/28-multitenancy.md — draft target design and current multi-session limitation.",
        "docs/29-engine-capability-matrix.md — feature parity and engine caveats.",
        ".env.example and message/webhook/template source modules — current defaults, limits and delivery behavior.",
    ]:
        list_item(doc, source, bullet_id)
    subhead(doc, "Official external source")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("WhatsApp Business Messaging Policy — opt-in, templates, customer-service window, data protection and enforcement (accessed 24 August 2026): ")
    set_font(r, size=10.2, color=BLACK)
    add_hyperlink(p, "whatsappbusiness.com/policy", "https://whatsappbusiness.com/policy/")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Meta WhatsApp Cloud API documentation — official alternative for regulated and commercial-scale use: ")
    set_font(r, size=10.2, color=BLACK)
    add_hyperlink(p, "developers.facebook.com/docs/whatsapp/cloud-api", "https://developers.facebook.com/docs/whatsapp/cloud-api")
    callout(doc, "Scope note", "This handbook explains the inspected local v0.23.2 codebase and operational recommendations as of 24 August 2026. WhatsApp policy and platform behavior can change. Re-check official policy before production or commercial launch. This is technical guidance, not legal advice.", kind="note")

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
